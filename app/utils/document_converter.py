"""
문서 변환 유틸리티

마크다운 형식의 공고문을 PDF, DOCX, HWP 형식으로 변환
Claude를 사용하여 마크다운을 PDF/DOCX에 최적화된 형식으로 변환
"""

from typing import Optional
import io
import re
import os
import tempfile
import subprocess
import logging
from pathlib import Path
from app.config import get_settings

logger = logging.getLogger(__name__)

try:
    import markdown
    from weasyprint import HTML, CSS
    from weasyprint.text.fonts import FontConfiguration
except ImportError:
    markdown = None
    HTML = None
    CSS = None
    FontConfiguration = None

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

try:
    from htmldocx import HtmlToDocx
except ImportError:
    HtmlToDocx = None


def markdown_to_pdf(markdown_content: str, output_path: Optional[str] = None) -> bytes:
    """
    마크다운을 PDF로 변환

    Args:
        markdown_content: 마크다운 형식의 텍스트
        output_path: 출력 파일 경로 (None이면 bytes 반환)

    Returns:
        PDF 파일 바이트 (output_path가 None인 경우)
    """
    if HTML is None:
        raise ImportError(
            "PDF 변환을 위해 다음 패키지가 필요합니다: "
            "pip install markdown weasyprint"
        )

    # 마크다운을 HTML로 변환
    html_content = markdown.markdown(
        markdown_content,
        extensions=['tables', 'fenced_code']
    )

    # HTML 스타일 추가 (공공문서 스타일)
    styled_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {{
                size: A4;
                margin: 2cm;
            }}
            body {{
                font-family: "맑은 고딕", "Malgun Gothic", sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                color: #000;
            }}
            h1 {{
                font-size: 18pt;
                font-weight: bold;
                margin-top: 20pt;
                margin-bottom: 10pt;
                text-align: center;
            }}
            h2 {{
                font-size: 14pt;
                font-weight: bold;
                margin-top: 15pt;
                margin-bottom: 8pt;
                border-bottom: 1px solid #ccc;
                padding-bottom: 3pt;
            }}
            h3 {{
                font-size: 12pt;
                font-weight: bold;
                margin-top: 10pt;
                margin-bottom: 5pt;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin: 10pt 0;
            }}
            th, td {{
                border: 1px solid #000;
                padding: 5pt;
                text-align: left;
            }}
            th {{
                background-color: #f0f0f0;
                font-weight: bold;
            }}
            p {{
                margin: 5pt 0;
            }}
            strong {{
                font-weight: bold;
            }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """

    # HTML을 PDF로 변환
    font_config = FontConfiguration()
    pdf_bytes = HTML(string=styled_html).write_pdf(font_config=font_config)

    if output_path:
        with open(output_path, 'wb') as f:
            f.write(pdf_bytes)
        return pdf_bytes
    else:
        return pdf_bytes


def markdown_to_docx(markdown_content: str, output_path: Optional[str] = None) -> bytes:
    """
    마크다운을 DOCX로 변환 (한글에서 열 수 있음)

    Args:
        markdown_content: 마크다운 형식의 텍스트
        output_path: 출력 파일 경로 (None이면 bytes 반환)

    Returns:
        DOCX 파일 바이트 (output_path가 None인 경우)
    """
    if Document is None:
        raise ImportError(
            "DOCX 변환을 위해 다음 패키지가 필요합니다: "
            "pip install python-docx"
        )

    # DOCX 문서 생성
    doc = Document()

    # 마크다운 파싱 (간단한 구현)
    lines = markdown_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # 제목 처리
        if line.startswith('# '):
            # H1
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            # H2
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            # H3
            doc.add_heading(line[4:], level=3)
        elif line.startswith('---'):
            # 구분선 (빈 줄로 대체)
            doc.add_paragraph('')
        elif line.startswith('|'):
            # 테이블 처리
            table_data = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                if row and not all(cell.startswith('-') for cell in row):  # 헤더 구분선 제외
                    table_data.append(row)
                i += 1
            i -= 1  # 다음 반복에서 현재 줄 다시 처리

            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Light Grid Accent 1'

                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_data
                        # 첫 번째 행은 헤더로 스타일링
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True

        elif line.startswith('- ') or line.startswith('* '):
            # 리스트 항목
            list_text = line[2:]
            doc.add_paragraph(list_text, style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            # 볼드 텍스트
            text = line[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
        else:
            # 일반 문단
            # 인라인 마크다운 처리 (간단한 구현)
            text = line
            # **볼드** 처리
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # 볼드 제거 (일단)
            # {플레이스홀더} 제거 (혹시 남아있을 경우)
            text = re.sub(r'\{[^}]+\}', '', text)
            
            if text.strip():
                doc.add_paragraph(text)

        i += 1

    # 문서 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)

    # 바이트로 변환
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    result_bytes = doc_bytes.read()

    if output_path:
        with open(output_path, 'wb') as f:
            f.write(result_bytes)
        return result_bytes
    else:
        return result_bytes


def convert_markdown_with_anthropic(
    markdown_content: str,
    output_format: str = "pdf"
) -> str:
    """
    Anthropic API를 직접 사용하여 마크다운을 PDF/DOCX에 최적화된 형식으로 변환
    
    Args:
        markdown_content: 마크다운 형식의 텍스트
        output_format: 출력 형식 ("pdf", "docx")
    
    Returns:
        변환된 HTML 또는 구조화된 텍스트
    """
    try:
        import anthropic
        settings = get_settings()
        
        if not settings.anthropic_api_key:
            print(f"⚠️ ANTHROPIC_API_KEY가 설정되지 않았습니다. 기존 라이브러리 사용")
            return None
        
        client = anthropic.Anthropic(api_key=settings.anthropic_api_key)
        model_name = os.getenv("ANTHROPIC_MODEL", settings.anthropic_model)
        
        format_instruction = {
            "pdf": "PDF 형식에 최적화된 완전한 HTML 문서로 변환하세요. <!DOCTYPE html><html><head><meta charset='UTF-8'><style>@page {size: A4; margin: 2cm;} body {font-family: '맑은 고딕', 'Malgun Gothic', sans-serif; font-size: 11pt; line-height: 1.6;}</style></head><body>...</body></html> 형식으로 완전한 HTML을 출력하세요.",
            "docx": "DOCX 형식에 최적화된 구조화된 마크다운으로 변환하세요. 제목, 단락, 테이블 구조를 명확히 구분하세요."
        }
        
        prompt = f"""
다음 마크다운 문서를 {output_format.upper()} 형식에 최적화된 형식으로 변환하세요.

{format_instruction.get(output_format.lower(), "")}

마크다운 내용:
```markdown
{markdown_content}
```

변환 규칙:
1. 모든 내용을 정확히 보존하세요
2. 섹션 구조를 명확히 유지하세요
3. 테이블, 리스트, 강조 표시를 올바르게 변환하세요
4. 한국어 폰트와 스타일을 고려하세요

변환된 결과만 출력하세요 (설명 없이).
"""
        
        response = client.messages.create(
            model=model_name,
            max_tokens=16000,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )
        
        # 응답 처리 (document_parser.py와 동일한 방식)
        if response.content and len(response.content) > 0:
            result_text = response.content[0].text
            if result_text and result_text.strip():
                print(f"✅ Anthropic API로 {output_format.upper()} 변환 성공 ({len(result_text)}자)")
                return result_text
            else:
                print(f"⚠️ Anthropic API 응답 텍스트가 비어있습니다. 기존 라이브러리 사용")
        else:
            print(f"⚠️ Anthropic API 응답에 content가 없습니다. 기존 라이브러리 사용")
        
        return None
        
    except Exception as e:
        print(f"⚠️ Anthropic API 변환 실패: {e}. 기존 라이브러리 사용")
        return None


def _is_html(content: str) -> bool:
    """
    내용이 HTML인지 마크다운인지 판단
    
    Args:
        content: 텍스트 내용
    
    Returns:
        HTML이면 True, 마크다운이면 False
    """
    content_stripped = content.strip()
    return (
        content_stripped.startswith("<!DOCTYPE html>") or
        content_stripped.startswith("<html>") or
        content_stripped.startswith("<HTML>") or
        ("<body>" in content_stripped.lower() and "<head>" in content_stripped.lower())
    )


def convert_document(
    content: str,
    output_format: str = "pdf",
    output_path: Optional[str] = None,
    is_html: Optional[bool] = None
) -> bytes:
    """
    문서를 지정된 형식으로 변환 (HTML 또는 마크다운 지원)
    
    HTML인 경우: convert_html_document 사용
    마크다운인 경우: 기존 로직 사용 (Claude 우선)

    Args:
        content: HTML 또는 마크다운 형식의 텍스트
        output_format: 출력 형식 ("pdf", "docx", "hwp")
        output_path: 출력 파일 경로 (None이면 bytes 반환)
        is_html: HTML 여부 (None이면 자동 감지)

    Returns:
        변환된 파일 바이트

    Raises:
        ValueError: 지원하지 않는 형식
        ImportError: 필요한 패키지가 설치되지 않음
    """
    # HTML 여부 자동 감지
    if is_html is None:
        is_html = _is_html(content)
    
    # HTML인 경우 직접 변환
    if is_html:
        return convert_html_document(content, output_format, output_path)
    
    # 마크다운인 경우 기존 로직 사용
    markdown_content = content
    # Anthropic API로 먼저 변환 시도
    anthropic_result = convert_markdown_with_anthropic(markdown_content, output_format)
    
    if anthropic_result and output_format.lower() == "pdf":
        # Anthropic이 HTML을 반환했다면, 이를 PDF로 변환
        if HTML is None:
            raise ImportError("PDF 변환을 위해 weasyprint가 필요합니다: pip install weasyprint")
        
        # Anthropic이 이미 완전한 HTML을 반환했는지 확인
        if anthropic_result.strip().startswith("<!DOCTYPE html>") or anthropic_result.strip().startswith("<html>"):
            # 이미 완전한 HTML이면 그대로 사용
            styled_html = anthropic_result
        else:
            # HTML body만 있으면 전체 HTML 구조로 감싸기
            styled_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    @page {{
                        size: A4;
                        margin: 2cm;
                    }}
                    body {{
                        font-family: "맑은 고딕", "Malgun Gothic", sans-serif;
                        font-size: 11pt;
                        line-height: 1.6;
                        color: #000;
                    }}
                </style>
            </head>
            <body>
                {anthropic_result}
            </body>
            </html>
            """
        
        font_config = FontConfiguration()
        pdf_bytes = HTML(string=styled_html).write_pdf(font_config=font_config)
        
        if output_path:
            with open(output_path, 'wb') as f:
                f.write(pdf_bytes)
        return pdf_bytes
    
    elif anthropic_result and output_format.lower() == "docx":
        # Anthropic이 구조화된 텍스트를 반환했다면, 이를 DOCX로 변환
        if Document is None:
            raise ImportError("DOCX 변환을 위해 python-docx가 필요합니다: pip install python-docx")
        
        # Anthropic 결과를 파싱하여 DOCX 생성 (기존 markdown_to_docx 로직 활용)
        return markdown_to_docx(anthropic_result, output_path)
    
    # Anthropic 변환 실패 시 기존 라이브러리 사용
    print("📝 Anthropic API 변환 실패 또는 미사용. 기존 라이브러리 사용")
    if output_format.lower() == "pdf":
        return markdown_to_pdf(markdown_content, output_path)
    elif output_format.lower() == "docx":
        return markdown_to_docx(markdown_content, output_path)
    elif output_format.lower() == "hwp":
        # HWP 변환은 HTML을 통해서만 가능
        # 마크다운을 HTML로 변환 후 HWP로 변환
        if HTML is None:
            raise ImportError("HWP 변환을 위해 weasyprint가 필요합니다: pip install weasyprint")
        
        # 마크다운을 HTML로 변환
        html_content = markdown.markdown(
            markdown_content,
            extensions=['tables', 'fenced_code']
        )
        
        # HTML 구조로 감싸기
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{
                    font-family: "맑은 고딕", "Malgun Gothic", sans-serif;
                    font-size: 11pt;
                    line-height: 1.6;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        return html_to_hwp_with_libreoffice(full_html, output_path)
    else:
        raise ValueError(f"지원하지 않는 형식: {output_format}. 'pdf' 또는 'docx'를 사용하세요.")


# 편의 함수
def export_to_file(
    markdown_content: str,
    output_format: str,
    filename: str
) -> str:
    """
    마크다운을 파일로 내보내기

    Args:
        markdown_content: 마크다운 형식의 텍스트
        output_format: 출력 형식 ("pdf", "docx")
        filename: 출력 파일명 (확장자 포함)

    Returns:
        저장된 파일 경로
    """
    output_path = Path(filename)
    convert_document(markdown_content, output_format, str(output_path))
    return str(output_path)


def _find_libreoffice() -> Optional[str]:
    """
    LibreOffice 실행 파일 경로 찾기
    
    Returns:
        LibreOffice 실행 파일 경로 또는 None
    """
    soffice_paths = [
        "/usr/bin/soffice",            # Linux (Docker)
        "/opt/homebrew/bin/soffice",  # macOS Homebrew
        "/usr/local/bin/soffice",      # macOS
        "/Applications/LibreOffice.app/Contents/MacOS/soffice"  # macOS app
    ]
    
    for path in soffice_paths:
        if os.path.exists(path):
            return path
    
    return None


def html_to_pdf(html_content: str, output_path: Optional[str] = None) -> bytes:
    """
    HTML을 PDF로 변환 (파란색 스타일 유지)
    
    Args:
        html_content: HTML 형식의 텍스트
        output_path: 출력 파일 경로 (None이면 bytes 반환)
    
    Returns:
        PDF 파일 바이트
    """
    if HTML is None:
        raise ImportError("PDF 변환을 위해 weasyprint가 필요합니다: pip install weasyprint")
    
    # HTML이 완전한 문서인지 확인
    if not html_content.strip().startswith("<!DOCTYPE html>") and not html_content.strip().startswith("<html>"):
        # HTML body만 있으면 전체 구조로 감싸기
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @font-face {{
                    font-family: 'NotoSansKR';
                    src: url('/fonts/NotoSansKR-Regular.ttf') format('truetype');
                }}
                @page {{
                    size: A4;
                    margin: 2cm;
                }}
                body {{
                    font-family: 'NotoSansKR', "맑은 고딕", "Malgun Gothic", sans-serif;
                    font-size: 11pt;
                    line-height: 1.6;
                    color: #000;
                }}
                .modified, .extracted {{
                    color: #0066CC;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
    
    # HTML에 charset이 명시되어 있는지 확인하고, 없으면 추가
    html_lower = html_content.lower()
    if '<meta charset' not in html_lower and '<meta http-equiv="content-type"' not in html_lower:
        # head 태그 안에 charset 메타 태그 추가
        if '<head>' in html_content:
            html_content = html_content.replace('<head>', '<head>\n    <meta charset="UTF-8">', 1)
        elif '<HEAD>' in html_content:
            html_content = html_content.replace('<HEAD>', '<HEAD>\n    <meta charset="UTF-8">', 1)
    
    # HTML에 NotoSansKR 폰트가 없으면 추가 (한글 폰트 문제 해결)
    html_lower_check = html_content.lower()
    if 'notosanskr' not in html_lower_check:
        # <style> 태그 찾기
        style_pattern = r'<style[^>]*>'
        style_match = re.search(style_pattern, html_content, re.IGNORECASE)
        
        if style_match:
            # <style> 태그 바로 다음에 @font-face 추가
            style_end = style_match.end()
            font_face_css = """    @font-face {
        font-family: 'NotoSansKR';
        src: url('/fonts/NotoSansKR-Regular.ttf') format('truetype');
    }
"""
            html_content = html_content[:style_end] + '\n' + font_face_css + html_content[style_end:]
            
            # body 스타일에 NotoSansKR 폰트 추가 (더 정확한 패턴 사용)
            # body { ... } 패턴 찾기 (여러 줄 지원)
            body_pattern = r'body\s*\{[^}]*\}'
            body_match = re.search(body_pattern, html_content, re.IGNORECASE | re.DOTALL)
            if body_match:
                body_style = body_match.group(0)
                if 'NotoSansKR' not in body_style and 'notosanskr' not in body_style.lower():
                    # font-family가 있으면 앞에 추가
                    if 'font-family' in body_style:
                        # font-family: ... ; 패턴 찾아서 앞에 NotoSansKR 추가
                        body_style = re.sub(
                            r'(font-family\s*:\s*)([^;]+)',
                            r"\1'NotoSansKR', \2",
                            body_style,
                            flags=re.IGNORECASE
                        )
                    else:
                        # font-family가 없으면 추가 (body { 다음에)
                        body_style = re.sub(
                            r'(body\s*\{)',
                            r"\1\n        font-family: 'NotoSansKR', sans-serif;",
                            body_style,
                            flags=re.IGNORECASE
                        )
                    html_content = html_content[:body_match.start()] + body_style + html_content[body_match.end():]
        else:
            # <style> 태그가 없으면 <head> 안에 추가
            if '<head>' in html_content or '<HEAD>' in html_content:
                head_end = html_content.find('</head>')
                if head_end == -1:
                    head_end = html_content.find('</HEAD>')
                if head_end > 0:
                    style_block = """
    <style>
        @font-face {
            font-family: 'NotoSansKR';
            src: url('/fonts/NotoSansKR-Regular.ttf') format('truetype');
        }
        body {
            font-family: 'NotoSansKR', sans-serif;
        }
    </style>
"""
                    html_content = html_content[:head_end] + style_block + html_content[head_end:]
    
    # WeasyPrint에 UTF-8로 전달하여 인코딩 문제 방지
    font_config = FontConfiguration()
    
    # HTML에 charset이 확실히 있는지 확인하고 추가
    html_lower = html_content.lower()
    has_charset = '<meta charset' in html_lower or 'charset=' in html_lower
    
    if not has_charset:
        # head 태그 바로 다음에 charset 추가
        if '<head>' in html_content:
            html_content = html_content.replace('<head>', '<head>\n    <meta charset="UTF-8">', 1)
        elif '<HEAD>' in html_content:
            html_content = html_content.replace('<HEAD>', '<HEAD>\n    <meta charset="UTF-8">', 1)
        else:
            # head 태그가 없으면 추가
            if '<html>' in html_content:
                html_content = html_content.replace('<html>', '<html>\n<head>\n    <meta charset="UTF-8">\n</head>', 1)
            elif '<HTML>' in html_content:
                html_content = html_content.replace('<HTML>', '<HTML>\n<HEAD>\n    <meta charset="UTF-8">\n</HEAD>', 1)
    
    # WeasyPrint에 UTF-8 바이트로 전달 (인코딩 문제 해결)
    # HTML의 charset을 더 명확하게 지정 (HTML5 표준)
    html_lower = html_content.lower()
    if '<meta charset' not in html_lower:
        # HTML5 방식: <meta charset="UTF-8">를 <head> 바로 다음에 추가
        if '<head>' in html_content:
            # <head> 다음에 charset 메타 태그 추가
            head_pos = html_content.find('<head>')
            if head_pos >= 0:
                head_end = html_content.find('>', head_pos) + 1
                html_content = html_content[:head_end] + '\n  <meta charset="UTF-8">' + html_content[head_end:]
        elif '<HEAD>' in html_content:
            head_pos = html_content.find('<HEAD>')
            if head_pos >= 0:
                head_end = html_content.find('>', head_pos) + 1
                html_content = html_content[:head_end] + '\n  <meta charset="UTF-8">' + html_content[head_end:]
    
    # HTML을 UTF-8로 완전히 정규화 (인코딩 문제 해결)
    # 한글 문자가 포함된 경우를 대비하여 UTF-8로 명시적으로 인코딩/디코딩
    try:
        # UTF-8로 인코딩 후 다시 디코딩하여 완전히 정규화
        html_normalized = html_content.encode('utf-8', errors='strict').decode('utf-8', errors='strict')
        logger.debug("HTML UTF-8 정규화 완료")
    except Exception as norm_err:
        logger.warning(f"HTML 정규화 실패, 원본 사용: {str(norm_err)}")
        html_normalized = html_content
    
    # HTML5 DOCTYPE이 없으면 추가 (WeasyPrint가 HTML5로 인식하도록)
    if not html_normalized.strip().startswith('<!DOCTYPE'):
        if html_normalized.strip().startswith('<html'):
            html_normalized = '<!DOCTYPE html>\n' + html_normalized
            logger.debug("HTML5 DOCTYPE 추가")
    
    # charset 메타 태그를 HTML5 방식으로 명확하게 지정 (WeasyPrint가 인식하도록)
    # <head> 태그의 첫 번째 자식으로 charset 메타 태그를 배치
    html_lower = html_normalized.lower()
    has_charset_meta = '<meta charset' in html_lower or 'charset=' in html_lower
    
    if not has_charset_meta:
        # HTML5 방식: <meta charset="UTF-8">를 <head> 바로 다음에 추가
        if '<head>' in html_normalized:
            head_pos = html_normalized.find('<head>')
            if head_pos >= 0:
                head_end = html_normalized.find('>', head_pos) + 1
                html_normalized = html_normalized[:head_end] + '\n  <meta charset="UTF-8">' + html_normalized[head_end:]
                logger.debug("charset 메타 태그 추가 (<head> 다음)")
        elif '<HEAD>' in html_normalized:
            head_pos = html_normalized.find('<HEAD>')
            if head_pos >= 0:
                head_end = html_normalized.find('>', head_pos) + 1
                html_normalized = html_normalized[:head_end] + '\n  <meta charset="UTF-8">' + html_normalized[head_end:]
                logger.debug("charset 메타 태그 추가 (<HEAD> 다음)")
        else:
            # head 태그가 없으면 html 태그 다음에 head 추가
            if '<html>' in html_normalized:
                html_pos = html_normalized.find('<html>')
                html_end = html_normalized.find('>', html_pos) + 1
                html_normalized = html_normalized[:html_end] + '\n<head>\n  <meta charset="UTF-8">\n</head>' + html_normalized[html_end:]
                logger.debug("head 태그와 charset 메타 태그 추가")
    else:
        # charset 메타 태그가 있지만 올바른 위치에 있는지 확인
        # <head> 태그 바로 다음에 오도록 재배치
        if '<head>' in html_normalized or '<HEAD>' in html_normalized:
            head_tag = '<head>' if '<head>' in html_normalized else '<HEAD>'
            head_pos = html_normalized.find(head_tag)
            if head_pos >= 0:
                head_end = html_normalized.find('>', head_pos) + 1
                # charset 메타 태그를 찾아서 제거하고 <head> 바로 다음에 재배치
                charset_pattern = r'<meta\s+charset=["\']?UTF-8["\']?\s*/?>'
                html_normalized = re.sub(charset_pattern, '', html_normalized, flags=re.IGNORECASE)
                html_normalized = html_normalized[:head_end] + '\n  <meta charset="UTF-8">' + html_normalized[head_end:]
                logger.debug("charset 메타 태그 재배치")
    
    # 방법 1: 임시 파일을 UTF-8 바이너리 모드로 저장 (가장 확실)
    # UTF-8 바이트로 저장하면 WeasyPrint가 charset 메타 태그를 읽어서 올바른 인코딩으로 파싱
    tmp_file_path = None
    try:
        # HTML을 UTF-8 바이트로 변환
        html_bytes = html_normalized.encode('utf-8', errors='strict')
        logger.debug(f"HTML을 UTF-8 바이트로 변환 완료 (크기: {len(html_bytes)} bytes)")
        
        # 바이너리 모드로 임시 파일 생성
        with tempfile.NamedTemporaryFile(mode='wb', suffix='.html', delete=False) as tmp_file:
            tmp_file.write(html_bytes)
            tmp_file_path = tmp_file.name
            logger.debug(f"임시 파일 생성: {tmp_file_path}")
        
        # 파일 경로로 HTML 로드 (WeasyPrint가 파일의 charset 메타 태그를 읽어서 인코딩 결정)
        logger.info("WeasyPrint HTML 파싱 시작...")
        pdf_bytes = HTML(filename=tmp_file_path).write_pdf(font_config=font_config)
        logger.info("✅ 임시 파일(UTF-8 바이너리)로 PDF 변환 성공")
    except UnicodeEncodeError as e:
        logger.warning(f"임시 파일(UTF-8 바이너리) 방법 실패 (인코딩 오류): {str(e)}")
        # 방법 2: UTF-8 텍스트 모드로 저장
        try:
            with tempfile.NamedTemporaryFile(mode='w', encoding='utf-8', suffix='.html', delete=False) as tmp_file:
                tmp_file.write(html_normalized)
                tmp_file_path = tmp_file.name
                logger.debug(f"임시 파일 생성 (텍스트 모드): {tmp_file_path}")
            
            pdf_bytes = HTML(filename=tmp_file_path).write_pdf(font_config=font_config)
            logger.info("✅ 임시 파일(UTF-8 텍스트)로 PDF 변환 성공")
        except Exception as e2:
            logger.warning(f"임시 파일(바이너리) 방법 실패: {str(e2)}")
            # 방법 3: UTF-8 바이트를 BytesIO로 전달
            try:
                html_bytes = html_normalized.encode('utf-8')
                html_file_obj = io.BytesIO(html_bytes)
                pdf_bytes = HTML(file_obj=html_file_obj, base_url='.').write_pdf(font_config=font_config)
                logger.info("✅ BytesIO로 PDF 변환 성공")
            except Exception as e3:
                logger.warning(f"BytesIO 방법 실패: {str(e3)}")
                # 방법 4: 문자열로 직접 전달 (최종 fallback)
                try:
                    pdf_bytes = HTML(string=html_normalized).write_pdf(font_config=font_config)
                    logger.info("✅ 문자열 직접 전달로 PDF 변환 성공")
                except Exception as e4:
                    logger.error(f"모든 PDF 변환 방법 실패")
                    raise RuntimeError(f"PDF 변환 실패 (모든 방법 시도): UTF-8텍스트파일={str(e)}, 바이너리파일={str(e2)}, BytesIO={str(e3)}, 문자열={str(e4)}")
    except Exception as e:
        logger.warning(f"임시 파일(UTF-8 바이너리) 방법 실패: {str(e)}")
        logger.debug(f"오류 타입: {type(e).__name__}, 메시지: {str(e)}")
        # 위와 동일한 fallback 로직
        try:
            with tempfile.NamedTemporaryFile(mode='w', encoding='utf-8', suffix='.html', delete=False) as tmp_file:
                tmp_file.write(html_normalized)
                tmp_file_path = tmp_file.name
                logger.debug(f"임시 파일 생성 (텍스트 모드, fallback): {tmp_file_path}")
            
            pdf_bytes = HTML(filename=tmp_file_path).write_pdf(font_config=font_config)
            logger.info("✅ 임시 파일(UTF-8 텍스트)로 PDF 변환 성공")
        except Exception as e2:
            logger.warning(f"임시 파일(바이너리) 방법 실패: {str(e2)}")
            try:
                html_bytes = html_normalized.encode('utf-8')
                html_file_obj = io.BytesIO(html_bytes)
                pdf_bytes = HTML(file_obj=html_file_obj, base_url='.').write_pdf(font_config=font_config)
                logger.info("✅ BytesIO로 PDF 변환 성공")
            except Exception as e3:
                logger.warning(f"BytesIO 방법 실패: {str(e3)}")
                try:
                    pdf_bytes = HTML(string=html_normalized).write_pdf(font_config=font_config)
                    logger.info("✅ 문자열 직접 전달로 PDF 변환 성공")
                except Exception as e4:
                    logger.error(f"모든 PDF 변환 방법 실패")
                    raise RuntimeError(f"PDF 변환 실패 (모든 방법 시도): UTF-8텍스트파일={str(e)}, 바이너리파일={str(e2)}, BytesIO={str(e3)}, 문자열={str(e4)}")
    finally:
        # 임시 파일 삭제
        if tmp_file_path and os.path.exists(tmp_file_path):
            try:
                os.unlink(tmp_file_path)
            except:
                pass
    
    if output_path:
        with open(output_path, 'wb') as f:
            f.write(pdf_bytes)
    
    return pdf_bytes


def html_to_docx(html_content: str, output_path: Optional[str] = None) -> bytes:
    """
    HTML을 DOCX로 변환 (HtmlToDocx 사용, 인코딩 문제 해결)
    
    Args:
        html_content: HTML 형식의 텍스트
        output_path: 출력 파일 경로 (None이면 bytes 반환)
    
    Returns:
        DOCX 파일 바이트
    
    Raises:
        ImportError: HtmlToDocx가 설치되지 않음
        RuntimeError: 변환 실패
    """
    logger.info(f"html_to_docx 함수 호출됨 (HtmlToDocx={HtmlToDocx is not None})")
    
    if HtmlToDocx is None:
        # HtmlToDocx가 없으면 LibreOffice fallback
        logger.warning("HtmlToDocx가 설치되지 않음, LibreOffice 사용")
        return html_to_docx_with_libreoffice(html_content, output_path)
    
    try:
        logger.info("HtmlToDocx로 HTML → DOCX 변환 시도...")
        # HtmlToDocx로 HTML을 DOCX로 변환
        parser = HtmlToDocx()
        
        # HTML 문자열을 DOCX Document로 변환
        logger.debug(f"HTML 내용 길이: {len(html_content)}")
        doc = parser.parse_html_string(html_content)
        logger.debug("HtmlToDocx.parse_html_string() 성공")
        
        # 임시 파일을 사용하여 DOCX 저장 (Document.save()는 파일 경로만 받음)
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            tmp_path = tmp_file.name
        
        try:
            # 임시 파일에 저장
            doc.save(tmp_path)
            logger.debug(f"DOCX 임시 파일 저장: {tmp_path}")
            
            # 파일 읽기
            with open(tmp_path, 'rb') as f:
                docx_content = f.read()
            logger.debug(f"DOCX 파일 크기: {len(docx_content)} bytes")
            
            # 출력 경로가 있으면 복사
            if output_path:
                with open(output_path, 'wb') as f:
                    f.write(docx_content)
            
            logger.info("✅ HTML → DOCX 변환 성공 (HtmlToDocx)")
            return docx_content
        finally:
            # 임시 파일 삭제
            if os.path.exists(tmp_path):
                try:
                    os.unlink(tmp_path)
                except:
                    pass
        
    except Exception as e:
        logger.error(f"HtmlToDocx 변환 실패: {str(e)}")
        import traceback
        logger.error(f"HtmlToDocx 오류 상세: {traceback.format_exc()}")
        logger.warning("LibreOffice fallback 시도...")
        # 실패 시 LibreOffice fallback
        return html_to_docx_with_libreoffice(html_content, output_path)


def html_to_docx_with_libreoffice(html_content: str, output_path: Optional[str] = None) -> bytes:
    """
    HTML을 DOCX로 변환 (LibreOffice 사용, 파란색 스타일 유지)
    
    Args:
        html_content: HTML 형식의 텍스트
        output_path: 출력 파일 경로 (None이면 bytes 반환)
    
    Returns:
        DOCX 파일 바이트
    
    Raises:
        RuntimeError: LibreOffice가 설치되지 않았거나 변환 실패
    """
    soffice_path = _find_libreoffice()
    
    if not soffice_path:
        raise RuntimeError(
            "LibreOffice가 설치되지 않았습니다. "
            "Docker 환경에서는 Dockerfile에 LibreOffice 설치가 필요합니다."
        )
    
    # HTML이 완전한 문서인지 확인
    if not html_content.strip().startswith("<!DOCTYPE html>") and not html_content.strip().startswith("<html>"):
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{
                    font-family: "맑은 고딕", "Malgun Gothic", sans-serif;
                    font-size: 11pt;
                    line-height: 1.6;
                }}
                .modified, .extracted {{
                    color: #0066CC;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
    
    # 임시 디렉토리 생성
    with tempfile.TemporaryDirectory() as temp_dir:
        # HTML 파일 저장
        html_path = os.path.join(temp_dir, "input.html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        
        # LibreOffice로 DOCX 변환
        try:
            result = subprocess.run(
                [
                    soffice_path,
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", temp_dir,
                    html_path
                ],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode != 0:
                raise RuntimeError(
                    f"HTML → DOCX 변환 실패 (exit code {result.returncode}): {result.stderr or result.stdout}"
                )
            
            # 변환된 DOCX 파일 읽기
            docx_path = os.path.join(temp_dir, "input.docx")
            if not os.path.exists(docx_path):
                raise RuntimeError(
                    f"DOCX 파일이 생성되지 않았습니다. 생성된 파일: {os.listdir(temp_dir)}"
                )
            
            with open(docx_path, "rb") as f:
                docx_content = f.read()
            
            if output_path:
                with open(output_path, 'wb') as f:
                    f.write(docx_content)
            
            print(f"✅ HTML → DOCX 변환 성공 (LibreOffice)")
            return docx_content
            
        except subprocess.TimeoutExpired:
            raise RuntimeError("HTML → DOCX 변환 시간 초과 (60초)")
        except Exception as e:
            raise RuntimeError(f"HTML → DOCX 변환 중 오류: {str(e)}")


def docx_to_pdf(docx_content: bytes, output_path: Optional[str] = None) -> bytes:
    """
    DOCX를 PDF로 변환 (LibreOffice 사용)
    
    Args:
        docx_content: DOCX 파일 바이트
        output_path: 출력 파일 경로 (None이면 bytes 반환)
    
    Returns:
        PDF 파일 바이트
    
    Raises:
        RuntimeError: LibreOffice가 설치되지 않았거나 변환 실패
    """
    soffice_path = _find_libreoffice()
    
    if not soffice_path:
        raise RuntimeError(
            "LibreOffice가 설치되지 않았습니다. "
            "Docker 환경에서는 Dockerfile에 LibreOffice 설치가 필요합니다."
        )
    
    # 임시 디렉토리 생성
    with tempfile.TemporaryDirectory() as temp_dir:
        # DOCX 파일 저장
        docx_path = os.path.join(temp_dir, "input.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_content)
        
        # LibreOffice로 PDF 변환
        try:
            result = subprocess.run(
                [
                    soffice_path,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", temp_dir,
                    docx_path
                ],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode != 0:
                raise RuntimeError(
                    f"DOCX → PDF 변환 실패 (exit code {result.returncode}): {result.stderr or result.stdout}"
                )
            
            # 변환된 PDF 파일 읽기
            pdf_path = os.path.join(temp_dir, "input.pdf")
            if not os.path.exists(pdf_path):
                raise RuntimeError(
                    f"PDF 파일이 생성되지 않았습니다. 생성된 파일: {os.listdir(temp_dir)}"
                )
            
            with open(pdf_path, "rb") as f:
                pdf_content = f.read()
            
            if output_path:
                with open(output_path, 'wb') as f:
                    f.write(pdf_content)
            
            logger.info("✅ DOCX → PDF 변환 성공 (LibreOffice)")
            return pdf_content
            
        except subprocess.TimeoutExpired:
            raise RuntimeError("DOCX → PDF 변환 시간 초과 (60초)")
        except Exception as e:
            raise RuntimeError(f"DOCX → PDF 변환 중 오류: {str(e)}")


def html_to_hwp_with_libreoffice(html_content: str, output_path: Optional[str] = None) -> bytes:
    """
    HTML을 HWP로 변환 (LibreOffice 사용, 파란색 스타일 유지)
    
    LibreOffice는 HWP로 직접 변환을 지원하지 않을 수 있으므로,
    HTML → DOCX → HWP 경로를 시도합니다.
    
    Args:
        html_content: HTML 형식의 텍스트
        output_path: 출력 파일 경로 (None이면 bytes 반환)
    
    Returns:
        HWP 파일 바이트
    
    Raises:
        RuntimeError: LibreOffice가 설치되지 않았거나 변환 실패
    """
    soffice_path = _find_libreoffice()
    
    if not soffice_path:
        raise RuntimeError(
            "LibreOffice가 설치되지 않았습니다. "
            "Docker 환경에서는 Dockerfile에 LibreOffice 설치가 필요합니다."
        )
    
    # HTML이 완전한 문서인지 확인
    if not html_content.strip().startswith("<!DOCTYPE html>") and not html_content.strip().startswith("<html>"):
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{
                    font-family: "맑은 고딕", "Malgun Gothic", sans-serif;
                    font-size: 11pt;
                    line-height: 1.6;
                }}
                .modified, .extracted {{
                    color: #0066CC;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
    
    # 임시 디렉토리 생성
    with tempfile.TemporaryDirectory() as temp_dir:
        # HTML 파일 저장
        html_path = os.path.join(temp_dir, "input.html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        
        # 방법 1: HTML → HWP 직접 변환 시도
        hwp_path = os.path.join(temp_dir, "input.hwp")
        try:
            logger.info("HTML → HWP 직접 변환 시도...")
            result = subprocess.run(
                [
                    soffice_path,
                    "--headless",
                    "--convert-to", "hwp",
                    "--outdir", temp_dir,
                    html_path
                ],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode == 0 and os.path.exists(hwp_path):
                with open(hwp_path, "rb") as f:
                    hwp_content = f.read()
                
                if output_path:
                    with open(output_path, 'wb') as f:
                        f.write(hwp_content)
                
                logger.info("✅ HTML → HWP 직접 변환 성공")
                return hwp_content
            else:
                logger.warning(f"HTML → HWP 직접 변환 실패: {result.stderr or result.stdout}")
        except Exception as e:
            logger.warning(f"HTML → HWP 직접 변환 시도 실패: {str(e)}")
        
        # 방법 2: HTML → DOCX로 변환 (LibreOffice는 HWP 변환을 지원하지 않음)
        # HWP는 한글과컴퓨터의 독점 포맷이므로 LibreOffice로는 변환 불가
        # 대신 DOCX로 변환하여 반환 (한글에서 열 수 있음)
        try:
            logger.info("LibreOffice는 HWP 변환을 지원하지 않습니다. HTML → DOCX로 변환합니다...")
            # HTML → DOCX
            docx_path = os.path.join(temp_dir, "input.docx")
            result1 = subprocess.run(
                [
                    soffice_path,
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", temp_dir,
                    html_path
                ],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result1.returncode != 0 or not os.path.exists(docx_path):
                raise RuntimeError(
                    f"HTML → DOCX 변환 실패 (exit code {result1.returncode}): {result1.stderr or result1.stdout}"
                )
            
            logger.info("✅ HTML → DOCX 변환 성공 (HWP는 LibreOffice에서 지원하지 않으므로 DOCX 반환)")
            
            # DOCX 파일 반환 (HWP 대신)
            with open(docx_path, "rb") as f:
                docx_content = f.read()
            
            if output_path:
                # 출력 경로가 있으면 .hwp 확장자를 .docx로 변경
                if output_path.endswith('.hwp'):
                    output_path = output_path[:-4] + '.docx'
                with open(output_path, 'wb') as f:
                    f.write(docx_content)
            
            logger.warning("⚠️ HWP 변환은 LibreOffice에서 지원하지 않습니다. DOCX 파일을 반환합니다.")
            logger.warning("   한글(HWP)에서 DOCX 파일을 열어서 HWP로 저장할 수 있습니다.")
            return docx_content
                
        except subprocess.TimeoutExpired:
            raise RuntimeError("HTML → DOCX 변환 시간 초과 (60초)")
        except Exception as e:
            raise RuntimeError(
                f"HTML → DOCX 변환 중 오류: {str(e)}\n"
                f"참고: LibreOffice는 HWP 변환을 지원하지 않습니다. DOCX로 변환을 시도했습니다."
            )


def mark_modified_text_in_html(html_content: str, modified_texts: list, extracted_texts: list = None) -> str:
    """
    HTML에서 수정/추출된 텍스트를 파란색으로 마킹
    
    Args:
        html_content: HTML 형식의 텍스트
        modified_texts: 수정된 텍스트 목록
        extracted_texts: 자동 추출된 텍스트 목록 (선택)
    
    Returns:
        파란색으로 마킹된 HTML
    """
    if extracted_texts is None:
        extracted_texts = []
    
    result = html_content
    
    # 수정된 텍스트를 파란색으로 마킹
    for text in modified_texts:
        if text and text.strip():
            # HTML 특수문자 이스케이프
            escaped_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            # 파란색 span으로 감싸기
            marked_text = f'<span class="modified" style="color: #0066CC;">{escaped_text}</span>'
            # 텍스트를 찾아서 교체 (대소문자 구분 없이)
            result = re.sub(
                re.escape(text),
                marked_text,
                result,
                flags=re.IGNORECASE
            )
    
    # 자동 추출된 텍스트를 파란색으로 마킹
    for text in extracted_texts:
        if text and text.strip():
            # HTML 특수문자 이스케이프
            escaped_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            # 파란색 span으로 감싸기
            marked_text = f'<span class="extracted" style="color: #0066CC;">{escaped_text}</span>'
            # 텍스트를 찾아서 교체 (대소문자 구분 없이)
            result = re.sub(
                re.escape(text),
                marked_text,
                result,
                flags=re.IGNORECASE
            )
    
    return result


def convert_html_document(
    html_content: str,
    output_format: str = "pdf",
    output_path: Optional[str] = None
) -> bytes:
    """
    HTML 문서를 지정된 형식으로 변환 (PDF, DOCX, HWP)
    
    Args:
        html_content: HTML 형식의 텍스트
        output_format: 출력 형식 ("pdf", "docx", "hwp")
        output_path: 출력 파일 경로 (None이면 bytes 반환)
    
    Returns:
        변환된 파일 바이트
    
    Raises:
        ValueError: 지원하지 않는 형식
        ImportError: 필요한 패키지가 설치되지 않음
        RuntimeError: LibreOffice 변환 실패
    """
    output_format_lower = output_format.lower()
    
    if output_format_lower == "pdf":
        # PDF 변환: WeasyPrint 직접 사용 (더 빠르고 안정적)
        logger.info("PDF 변환: WeasyPrint 직접 사용")
        return html_to_pdf(html_content, output_path)
    elif output_format_lower == "docx":
        logger.info("convert_html_document: DOCX 변환 요청")
        return html_to_docx(html_content, output_path)
    elif output_format_lower == "hwp":
        return html_to_hwp_with_libreoffice(html_content, output_path)
    else:
        raise ValueError(f"지원하지 않는 형식: {output_format}. 'pdf', 'docx', 또는 'hwp'를 사용하세요.")


def hwp_to_pdf(hwp_content: bytes) -> bytes:
    """
    HWP 파일을 PDF로 변환 (LibreOffice 사용)

    Args:
        hwp_content: HWP 파일 바이트

    Returns:
        PDF 파일 바이트

    Raises:
        RuntimeError: LibreOffice가 설치되지 않았거나 변환 실패
    """
    # LibreOffice 경로 확인
    soffice_path = _find_libreoffice()
    
    if not soffice_path:
        raise RuntimeError(
            "LibreOffice가 설치되지 않았습니다. "
            "Docker 환경에서는 Dockerfile에 LibreOffice 설치가 필요합니다."
        )

    # 임시 디렉토리 생성
    with tempfile.TemporaryDirectory() as temp_dir:
        # HWP 파일 저장
        hwp_path = os.path.join(temp_dir, "input.hwp")
        with open(hwp_path, "wb") as f:
            f.write(hwp_content)

        # LibreOffice로 PDF 변환
        try:
            result = subprocess.run(
                [
                    soffice_path,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", temp_dir,
                    hwp_path
                ],
                capture_output=True,
                text=True,
                timeout=60  # 60초 타임아웃
            )

            # 디버깅: 출력 확인
            print(f"🔍 LibreOffice 실행 결과:")
            print(f"  Return code: {result.returncode}")
            print(f"  STDOUT: {result.stdout}")
            print(f"  STDERR: {result.stderr}")

            # 생성된 파일 목록 확인
            generated_files = os.listdir(temp_dir)
            print(f"  생성된 파일들: {generated_files}")

            if result.returncode != 0:
                raise RuntimeError(
                    f"HWP → PDF 변환 실패 (exit code {result.returncode}): {result.stderr or result.stdout}"
                )

            # 변환된 PDF 파일 읽기
            pdf_path = os.path.join(temp_dir, "input.pdf")
            if not os.path.exists(pdf_path):
                raise RuntimeError(
                    f"PDF 파일이 생성되지 않았습니다. 생성된 파일: {generated_files}"
                )

            with open(pdf_path, "rb") as f:
                pdf_content = f.read()

            return pdf_content

        except subprocess.TimeoutExpired:
            raise RuntimeError("HWP → PDF 변환 시간 초과 (60초)")
        except Exception as e:
            raise RuntimeError(f"HWP → PDF 변환 중 오류: {str(e)}")


