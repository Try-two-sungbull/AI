"""
문서 파싱 유틸리티

PDF, DOCX, HWP 지원
- HWP 5.0 이전 버전 (OLE 기반)
- HWP 5.0+ 버전 (ZIP 기반)
"""

from pypdf import PdfReader
import pdfplumber
from docx import Document
import olefile
import io
import struct
from typing import Union
import chardet
import logging

logger = logging.getLogger(__name__)


def decode_text_with_fallback(data: bytes) -> str:
    """
    다양한 인코딩을 시도하여 텍스트 디코딩

    Args:
        data: 바이트 데이터

    Returns:
        디코딩된 텍스트
    """
    # 1. chardet으로 인코딩 감지
    try:
        detected = chardet.detect(data)
        if detected['encoding'] and detected['confidence'] > 0.7:
            return data.decode(detected['encoding'], errors='replace')
    except:
        pass

    # 2. 일반적인 인코딩 순차 시도
    encodings = [
        'utf-8',
        'cp949',  # 한국어 Windows
        'euc-kr',  # 한국어 Unix/Linux
        'utf-16',
        'utf-16le',
        'utf-16be',
        'latin-1'  # 최후의 수단 (거의 모든 바이트 시퀀스 허용)
    ]

    for encoding in encodings:
        try:
            decoded = data.decode(encoding)
            # 유효한 텍스트인지 간단히 확인
            if len(decoded.strip()) > 0:
                return decoded
        except (UnicodeDecodeError, LookupError):
            continue

    # 3. 최종 폴백: replace 모드로 UTF-8 디코딩
    return data.decode('utf-8', errors='replace')


def parse_document(file_content: bytes, filename: str) -> str:
    """
    업로드된 문서에서 텍스트 추출

    Args:
        file_content: 파일 바이트 내용
        filename: 파일 이름

    Returns:
        추출된 텍스트

    Raises:
        ValueError: 지원하지 않는 파일 형식
    """
    file_extension = filename.lower().split('.')[-1]

    if file_extension == 'pdf':
        return parse_pdf(file_content)
    elif file_extension in ['docx', 'doc']:
        return parse_docx(file_content)
    elif file_extension == 'hwp':
        # HWP 파일 직접 파싱
        # 참고: LibreOffice는 HWP 파일을 직접 읽을 수 없으므로 변환 불가
        logger.info("=" * 50)
        logger.info("📄 HWP 파일 직접 파싱 시작...")
        logger.info(f"   파일명: {filename}")
        logger.info(f"   파일 크기: {len(file_content)} bytes")
        logger.info("   ⚠️ 참고: HWP는 한글과컴퓨터 독점 포맷으로 PDF 자동 변환이 불가능합니다.")
        logger.info("   💡 더 나은 결과를 원하시면 HWP를 PDF로 변환 후 업로드해주세요.")
        logger.info("=" * 50)
        return parse_hwp(file_content)
    elif file_extension == 'txt':
        return decode_text_with_fallback(file_content)
    else:
        raise ValueError(f"지원하지 않는 파일 형식: {file_extension}")


def parse_pdf(file_content: bytes) -> str:
    """
    PDF 파일에서 텍스트 추출
    
    pypdf를 먼저 시도하고, 실패 시 pdfplumber를 fallback으로 사용

    Args:
        file_content: PDF 파일 바이트

    Returns:
        추출된 텍스트
    """
    # 방법 1: pypdf 시도
    try:
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        # 암호화된 PDF 체크
        if reader.is_encrypted:
            try:
                reader.decrypt("")  # 빈 비밀번호 시도
            except:
                raise ValueError("PDF가 암호화되어 있습니다. 비밀번호를 제거한 후 다시 시도해주세요.")

        text_parts = []
        total_pages = len(reader.pages)
        
        for page_num, page in enumerate(reader.pages, 1):
            try:
                text = page.extract_text()
                if text:
                    # PDF 추출 텍스트에 인코딩 문제가 있을 수 있음
                    # 깨진 문자 정리
                    cleaned_text = text.encode('utf-8', errors='ignore').decode('utf-8')
                    if cleaned_text.strip():
                        text_parts.append(cleaned_text)
            except Exception as page_error:
                # 개별 페이지 오류는 로그하고 계속
                logger.warning(f"페이지 {page_num}/{total_pages}에서 텍스트 추출 실패: {str(page_error)}")
                continue

        full_text = '\n\n'.join(text_parts)

        if full_text.strip():
            return clean_text(full_text)
        else:
            logger.warning("pypdf로 텍스트를 추출할 수 없어 pdfplumber를 시도합니다.")
            
    except Exception as e:
        logger.warning(f"pypdf 파싱 실패 ({str(e)}), pdfplumber를 시도합니다.")

    # 방법 2: pdfplumber fallback
    try:
        pdf_file = io.BytesIO(file_content)
        text_parts = []
        
        with pdfplumber.open(pdf_file) as pdf:
            total_pages = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(text.strip())
                except Exception as page_error:
                    logger.warning(f"pdfplumber로 페이지 {page_num}/{total_pages} 추출 실패: {str(page_error)}")
                    continue

        full_text = '\n\n'.join(text_parts)

        if full_text.strip():
            return clean_text(full_text)
            
    except Exception as e:
        logger.error(f"pdfplumber 파싱도 실패: {str(e)}")

    # 방법 3: Claude Vision API 사용 (최종 fallback)
    try:
        logger.info("pypdf와 pdfplumber 모두 실패, Claude Vision API를 사용하여 PDF에서 텍스트 추출을 시도합니다.")
        return parse_pdf_with_claude_vision(file_content)
    except Exception as e:
        logger.error(f"Claude Vision API 파싱도 실패: {str(e)}")
        raise ValueError(
            f"PDF에서 텍스트를 추출할 수 없습니다. 모든 방법 시도 실패: {str(e)}"
        )


def parse_docx(file_content: bytes) -> str:
    """
    DOCX 파일에서 텍스트 추출

    Args:
        file_content: DOCX 파일 바이트

    Returns:
        추출된 텍스트
    """
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)

        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # 표(table)에서도 텍스트 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        full_text = '\n'.join(text_parts)

        if not full_text.strip():
            raise ValueError("DOCX에서 텍스트를 추출할 수 없습니다")

        return clean_text(full_text)

    except Exception as e:
        raise ValueError(f"DOCX 파싱 실패: {str(e)}")


def parse_hwp(file_content: bytes) -> str:
    """
    HWP 파일에서 텍스트 추출 (개선된 버전)

    HWP 5.0 이전 버전(OLE 기반) 및 5.0+ 버전 지원
    실패 시 Claude Vision API를 사용하여 이미지로 변환 후 텍스트 추출

    Args:
        file_content: HWP 파일 바이트

    Returns:
        추출된 텍스트
    """
    try:
        hwp_file = io.BytesIO(file_content)

        # HWP 5.0+ (ZIP 기반) 확인
        if file_content[:2] == b'PK':
            try:
                return parse_hwp_50_plus(file_content)
            except Exception as e:
                logger.warning(f"HWP 5.0+ 파싱 실패, Claude Vision API 시도: {str(e)}")
                # Claude Vision API fallback
                return parse_hwp_with_claude_vision(file_content)

        # HWP 5.0 이전 버전 (OLE 기반)
        if not olefile.isOleFile(hwp_file):
            raise ValueError("지원하지 않는 HWP 파일 형식입니다")

        ole = olefile.OleFileIO(hwp_file)

        # HWP 파일의 텍스트 스트림 찾기
        text_parts = []

        # BodyText 섹션에서 텍스트 추출
        for dir_entry in ole.listdir():
            if 'BodyText' in dir_entry or 'Section' in str(dir_entry):
                try:
                    stream = ole.openstream(dir_entry)
                    data = stream.read()

                    # 개선된 텍스트 추출
                    # 다양한 인코딩 시도
                    try:
                        decoded_text = decode_text_with_fallback(data)
                        # 제어 문자 제거 (하지만 줄바꿈은 유지)
                        cleaned = ''.join(
                            char for char in decoded_text 
                            if char.isprintable() or char in ['\n', '\r', '\t', ' ']
                        )
                        # 연속된 공백 정리
                        import re
                        cleaned = re.sub(r'[ \t]+', ' ', cleaned)  # 공백 정리
                        cleaned = re.sub(r'\n\s*\n\s*\n+', '\n\n', cleaned)  # 연속된 줄바꿈 정리
                        
                        if cleaned.strip():
                            text_parts.append(cleaned.strip())
                    except:
                        pass
                except:
                    continue

        ole.close()

        full_text = '\n'.join(text_parts)

        if not full_text.strip():
            logger.warning("HWP OLE 파싱으로 텍스트를 추출할 수 없어 Claude Vision API를 시도합니다.")
            # Claude Vision API fallback
            return parse_hwp_with_claude_vision(file_content)

        return clean_text(full_text)

    except Exception as e:
        logger.warning(f"HWP 파싱 실패, Claude Vision API 시도: {str(e)}")
        # Claude Vision API fallback
        try:
            return parse_hwp_with_claude_vision(file_content)
        except Exception as vision_error:
            raise ValueError(
                f"HWP 파싱 실패 (모든 방법 시도): 기본 파싱={str(e)}, Claude Vision={str(vision_error)}"
            )


def parse_hwp_50_plus(file_content: bytes) -> str:
    """
    HWP 5.0+ (ZIP 기반) 파일에서 텍스트 추출 (개선된 버전)

    Args:
        file_content: HWP 파일 바이트

    Returns:
        추출된 텍스트
    """
    import zipfile
    import xml.etree.ElementTree as ET

    try:
        hwp_file = io.BytesIO(file_content)
        with zipfile.ZipFile(hwp_file, 'r') as zf:
            text_parts = []
            section_order = []

            # section*.xml 파일들을 순서대로 정렬
            section_files = [f for f in zf.namelist() 
                           if f.startswith('Contents/section') and f.endswith('.xml')]
            section_files.sort(key=lambda x: int(x.split('section')[1].split('.')[0]) if x.split('section')[1].split('.')[0].isdigit() else 999)

            # 각 섹션에서 텍스트 추출
            for filename in section_files:
                try:
                    xml_content = zf.read(filename)
                    # 인코딩 문제 해결: UTF-8로 디코딩 시도
                    try:
                        xml_text = xml_content.decode('utf-8')
                    except UnicodeDecodeError:
                        # UTF-8 실패 시 cp949 시도
                        try:
                            xml_text = xml_content.decode('cp949', errors='ignore')
                        except:
                            xml_text = xml_content.decode('utf-8', errors='replace')
                    
                    root = ET.fromstring(xml_text)
                    section_text = []
                    
                    # 더 정확한 텍스트 추출: 특정 태그 우선 처리
                    # t 태그 (텍스트), tc 태그 (표 셀), li 태그 (목록 항목) 등
                    for elem in root.iter():
                        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                        
                        # 텍스트 노드 처리
                        if elem.text and elem.text.strip():
                            text = elem.text.strip()
                            # 중복 제거 및 의미있는 텍스트만 추가
                            if text and len(text) > 0:
                                section_text.append(text)
                        
                        # 표 셀 처리 (tc 태그)
                        if tag == 'tc':
                            cell_texts = []
                            for sub_elem in elem.iter():
                                if sub_elem.text and sub_elem.text.strip():
                                    cell_texts.append(sub_elem.text.strip())
                            if cell_texts:
                                section_text.append(' | '.join(cell_texts))
                    
                    if section_text:
                        # 섹션 구분을 위해 빈 줄 추가
                        if text_parts:
                            text_parts.append('')
                        text_parts.extend(section_text)
                        
                except Exception as e:
                    logger.warning(f"HWP 섹션 {filename} 파싱 실패: {str(e)}")
                    continue

            full_text = '\n'.join(text_parts)

            if not full_text.strip():
                raise ValueError("HWP 5.0+ 파일에서 텍스트를 추출할 수 없습니다")

            return clean_text(full_text)

    except Exception as e:
        raise ValueError(f"HWP 5.0+ 파싱 실패: {str(e)}")


def parse_pdf_with_claude_vision(file_content: bytes) -> str:
    """
    Claude Vision API를 사용하여 PDF에서 텍스트 추출 (최종 fallback)
    
    PDF를 이미지로 변환한 후 Claude에게 텍스트 추출을 요청합니다.
    
    Args:
        file_content: PDF 파일 바이트
    
    Returns:
        추출된 텍스트
    """
    import base64
    from app.config import get_settings
    
    settings = get_settings()
    
    # Claude API 키 확인
    if not settings.anthropic_api_key:
        raise ValueError("Claude Vision API를 사용하려면 ANTHROPIC_API_KEY가 필요합니다.")
    
    try:
        # PDF를 이미지로 변환 시도
        try:
            from pdf2image import convert_from_bytes
            images = convert_from_bytes(file_content, dpi=200, fmt='png')
        except ImportError:
            # pdf2image가 없으면 PyMuPDF 사용 시도
            try:
                import fitz  # PyMuPDF
                pdf_doc = fitz.open(stream=file_content, filetype="pdf")
                images = []
                for page_num in range(len(pdf_doc)):
                    page = pdf_doc[page_num]
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better quality
                    img_data = pix.tobytes("png")
                    from PIL import Image
                    import io
                    images.append(Image.open(io.BytesIO(img_data)))
                pdf_doc.close()
            except ImportError:
                raise ValueError("PDF를 이미지로 변환하려면 pdf2image 또는 PyMuPDF가 필요합니다.")
    
        # Claude API 직접 호출
        import anthropic
        import os
        client = anthropic.Anthropic(api_key=settings.anthropic_api_key)
        
        # 환경변수 또는 설정에서 모델명 가져오기
        model_name = os.getenv("ANTHROPIC_MODEL", settings.anthropic_model)
        logger.info(f"Claude Vision API 모델 사용: {model_name}")
        
        text_parts = []
        total_pages = len(images)
        
        # 페이지가 많으면 배치 처리 (비용 절감)
        max_pages_per_request = 5  # Claude는 최대 20개 이미지까지 처리 가능하지만, 비용 절감을 위해 5개씩
        
        for batch_start in range(0, total_pages, max_pages_per_request):
            batch_end = min(batch_start + max_pages_per_request, total_pages)
            batch_images = images[batch_start:batch_end]
            
            # 이미지를 base64로 인코딩
            image_contents = []
            for img in batch_images:
                import io
                img_bytes = io.BytesIO()
                img.save(img_bytes, format='PNG')
                img_bytes.seek(0)
                img_base64 = base64.b64encode(img_bytes.read()).decode('utf-8')
                image_contents.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": "image/png",
                        "data": img_base64
                    }
                })
            
            # Claude에게 텍스트 추출 요청
            message_content = [
                {
                    "type": "text",
                    "text": "이 PDF 이미지에서 모든 텍스트를 추출해주세요. 문서 구조와 형식을 최대한 유지하면서 텍스트만 반환해주세요. 표나 목록이 있으면 그 구조도 유지해주세요."
                }
            ]
            message_content.extend(image_contents)
            
            try:
                response = client.messages.create(
                    model=model_name,  # 환경변수 또는 설정에서 가져온 모델명 사용
                    max_tokens=8192,  # 더 많은 토큰으로 증가
                    messages=[
                        {
                            "role": "user",
                            "content": message_content
                        }
                    ]
                )
                
                # 응답 처리
                if response.content and len(response.content) > 0:
                    extracted_text = response.content[0].text
                    if extracted_text and extracted_text.strip():
                        text_parts.append(extracted_text)
                        logger.info(f"Claude Vision API로 페이지 {batch_start+1}-{batch_end} 텍스트 추출 성공 ({len(extracted_text)}자)")
                    else:
                        logger.warning(f"Claude Vision API로 페이지 {batch_start+1}-{batch_end} 추출된 텍스트가 비어있습니다.")
                else:
                    logger.warning(f"Claude Vision API 응답에 content가 없습니다.")
                    
            except Exception as e:
                error_detail = str(e)
                # 404 에러면 모델명 문제일 수 있음, 다른 모델로 fallback 시도
                if "404" in error_detail or "not_found" in error_detail.lower():
                    logger.warning(f"모델 {model_name}을 찾을 수 없습니다. claude-3-5-sonnet-20240620으로 fallback 시도합니다.")
                    try:
                        response = client.messages.create(
                            model="claude-3-5-sonnet-20240620",  # Fallback 모델
                            max_tokens=8192,
                            messages=[
                                {
                                    "role": "user",
                                    "content": message_content
                                }
                            ]
                        )
                        if response.content and len(response.content) > 0:
                            extracted_text = response.content[0].text
                            if extracted_text and extracted_text.strip():
                                text_parts.append(extracted_text)
                                logger.info(f"Fallback 모델로 페이지 {batch_start+1}-{batch_end} 텍스트 추출 성공 ({len(extracted_text)}자)")
                                continue
                    except Exception as fallback_error:
                        logger.error(f"Fallback 모델도 실패: {str(fallback_error)}")
                
                logger.error(f"Claude Vision API로 페이지 {batch_start+1}-{batch_end} 추출 실패: {error_detail}")
                # 전체 실패가 아니면 계속 시도
                continue
        
        full_text = '\n\n'.join(text_parts)
        
        if full_text.strip():
            logger.info(f"Claude Vision API로 총 {total_pages}페이지 중 {len(text_parts)}개 배치 추출 성공")
            return clean_text(full_text)
        else:
            raise ValueError(
                f"Claude Vision API로도 텍스트를 추출할 수 없습니다. "
                f"({total_pages}페이지 시도, {len(text_parts)}개 배치 성공)"
            )
            
    except ValueError as e:
        # ValueError는 그대로 전달
        raise
    except Exception as e:
        error_detail = str(e)
        logger.error(f"Claude Vision API 사용 중 예상치 못한 오류: {error_detail}", exc_info=True)
        raise ValueError(f"Claude Vision API 사용 실패: {error_detail}")


def parse_hwp_with_claude_vision(file_content: bytes) -> str:
    """
    Claude Vision API를 사용하여 HWP에서 텍스트 추출 (최종 fallback)
    
    HWP를 PDF로 변환한 후 이미지로 변환하여 Claude에게 텍스트 추출을 요청합니다.
    
    Args:
        file_content: HWP 파일 바이트
    
    Returns:
        추출된 텍스트
    """
    import base64
    import tempfile
    import subprocess
    from app.config import get_settings
    
    settings = get_settings()
    
    # Claude API 키 확인
    if not settings.anthropic_api_key:
        raise ValueError("Claude Vision API를 사용하려면 ANTHROPIC_API_KEY가 필요합니다.")
    
    try:
        # HWP를 PDF로 변환 (LibreOffice 사용)
        logger.info("HWP → PDF 변환 시도 (Claude Vision API용)...")
        
        # LibreOffice 경로 찾기
        soffice_paths = [
            "/usr/bin/soffice",
            "/opt/homebrew/bin/soffice",
            "/usr/local/bin/soffice",
        ]
        soffice_path = None
        for path in soffice_paths:
            import os
            if os.path.exists(path):
                soffice_path = path
                break
        
        if not soffice_path:
            raise ValueError("LibreOffice가 설치되지 않았습니다. HWP → PDF 변환을 위해 필요합니다.")
        
        # 임시 디렉토리 생성
        with tempfile.TemporaryDirectory() as temp_dir:
            # HWP 파일 저장
            hwp_path = os.path.join(temp_dir, "input.hwp")
            with open(hwp_path, "wb") as f:
                f.write(file_content)
            
            # LibreOffice로 PDF 변환
            result = subprocess.run(
                [
                    soffice_path,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", temp_dir,
                    hwp_path
                ],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode != 0:
                raise RuntimeError(f"HWP → PDF 변환 실패: {result.stderr or result.stdout}")
            
            # 변환된 PDF 파일 읽기
            pdf_path = os.path.join(temp_dir, "input.pdf")
            if not os.path.exists(pdf_path):
                raise RuntimeError("PDF 파일이 생성되지 않았습니다.")
            
            with open(pdf_path, "rb") as f:
                pdf_content = f.read()
            
            logger.info("✅ HWP → PDF 변환 성공, Claude Vision API로 텍스트 추출 시작...")
            
            # PDF를 Claude Vision API로 처리
            return parse_pdf_with_claude_vision(pdf_content)
            
    except Exception as e:
        error_detail = str(e)
        logger.error(f"HWP Claude Vision API 사용 실패: {error_detail}", exc_info=True)
        raise ValueError(f"HWP Claude Vision API 사용 실패: {error_detail}")


def clean_text(text: str) -> str:
    """
    추출된 텍스트 정리

    - 과도한 공백 제거
    - 줄바꿈 정리
    - 깨진 문자 제거

    Args:
        text: 원본 텍스트

    Returns:
        정리된 텍스트
    """
    import re

    # 유니코드 정규화 (NFC)
    import unicodedata
    text = unicodedata.normalize('NFC', text)

    # 제어 문자 제거 (줄바꿈, 탭 제외)
    text = ''.join(char for char in text if char.isprintable() or char in ['\n', '\r', '\t', ' '])

    # 연속된 공백을 하나로 (단, 줄바꿈은 유지)
    lines = text.split('\n')
    cleaned_lines = [' '.join(line.split()) for line in lines]
    text = '\n'.join(cleaned_lines)

    # 연속된 줄바꿈을 최대 2개로 제한
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()
